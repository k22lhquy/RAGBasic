# src/indexing.py
import os
import json
import hashlib
from src.embedding import get_embedding_model
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_chroma import Chroma
from langchain_core.documents import Document

from src.config import (
    EMBEDDING_MODEL, VECTORSTORE_DIR, COLLECTION_NAME,
    CHUNK_SIZE, CHUNK_OVERLAP
)

# File tracking những file đã được index
INDEX_REGISTRY = "data/indexed_files.json"


# ── Registry helpers ───────────────────────────────────────────────────────────

def _load_registry() -> dict:
    """Đọc registry {filename: md5}"""
    if not os.path.exists(INDEX_REGISTRY):
        return {}
    try:
        with open(INDEX_REGISTRY, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _save_registry(registry: dict):
    os.makedirs(os.path.dirname(INDEX_REGISTRY), exist_ok=True)
    with open(INDEX_REGISTRY, "w", encoding="utf-8") as f:
        json.dump(registry, f, ensure_ascii=False, indent=2)


def _file_md5(filepath: str) -> str:
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# ── Loaders ───────────────────────────────────────────────────────────────────

def _load_txt(fpath: str) -> Document | None:
    for encoding in ["utf-8", "utf-8-sig", "cp1258", "cp1252", "latin-1"]:
        try:
            with open(fpath, "r", encoding=encoding) as f:
                text = f.read()
            print(f"  ✅ TXT [{encoding}]: {fpath}")
            return Document(page_content=text, metadata={"source": fpath})
        except Exception:
            continue
    print(f"  ❌ Không đọc được: {fpath}")
    return None


def _load_pdf(fpath: str) -> list[Document]:
    try:
        loader = PyPDFLoader(fpath)
        docs = loader.load()
        print(f"  ✅ PDF: {fpath} ({len(docs)} trang)")
        return docs
    except Exception as e:
        print(f"  ❌ Lỗi PDF {fpath}: {e}")
        return []


def _load_docx(fpath: str) -> Document | None:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(fpath)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        print(f"  ✅ DOCX: {fpath}")
        return Document(page_content=text, metadata={"source": fpath})
    except Exception as e:
        print(f"  ❌ Lỗi DOCX {fpath}: {e}")
        return None


def load_single_file(fpath: str) -> list[Document]:
    """Load 1 file bất kỳ (TXT / PDF / DOCX) thành list[Document]"""
    ext = fpath.lower()
    if ext.endswith(".txt"):
        doc = _load_txt(fpath)
        return [doc] if doc else []
    elif ext.endswith(".pdf"):
        return _load_pdf(fpath)
    elif ext.endswith(".docx") or ext.endswith(".doc"):
        doc = _load_docx(fpath)
        return [doc] if doc else []
    else:
        print(f"  ⚠️  Bỏ qua file không hỗ trợ: {fpath}")
        return []


def load_documents(data_dir: str = "data") -> list[Document]:
    """Load toàn bộ file trong thư mục"""
    all_files = []
    for root, dirs, files in os.walk(data_dir):
        for f in files:
            # Bỏ qua file JSON registry
            if f.endswith(".json"):
                continue
            all_files.append(os.path.join(root, f))

    print(f"📁 File tìm thấy: {all_files}")
    documents = []
    for fpath in all_files:
        documents.extend(load_single_file(fpath))

    print(f"✅ Tổng cộng: {len(documents)} document")
    return documents


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunks = splitter.split_documents(documents)
    print(f"✅ Chunked: {len(chunks)} chunks")
    return chunks


# ── VectorStore helpers ───────────────────────────────────────────────────────

def build_vectorstore(chunks: list[Document], embeddings) -> Chroma:
    print("⏳ Đang embed và lưu vào VectorDB...")
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=VECTORSTORE_DIR,
        collection_name=COLLECTION_NAME
    )
    print(f"✅ Đã lưu {len(chunks)} chunks")
    return vectorstore


def load_vectorstore(embeddings) -> Chroma:
    return Chroma(
        persist_directory=VECTORSTORE_DIR,
        embedding_function=embeddings,
        collection_name=COLLECTION_NAME
    )


# ── Incremental API ───────────────────────────────────────────────────────────

def add_document(filepath: str, embeddings=None) -> dict:
    """
    Index 1 file mới vào vectorstore hiện có (không xóa file cũ).
    Trả về {"status": "added"|"skipped"|"error", "filename": ..., "chunks": ...}
    """
    filename = os.path.basename(filepath)
    registry = _load_registry()

    # Kiểm tra file đã index chưa (so hash)
    try:
        current_md5 = _file_md5(filepath)
    except Exception as e:
        return {"status": "error", "filename": filename, "message": str(e)}

    if filename in registry and registry[filename] == current_md5:
        print(f"   ⏭️  Bỏ qua (chưa thay đổi): {filename}")
        return {"status": "skipped", "filename": filename, "message": "File chưa thay đổi"}

    # Load & chunk
    docs = load_single_file(filepath)
    if not docs:
        return {"status": "error", "filename": filename, "message": "Không đọc được file"}

    chunks = chunk_documents(docs)

    # Lấy/tạo vectorstore
    if embeddings is None:
        embeddings = get_embedding_model()

    if os.path.exists(VECTORSTORE_DIR) and os.listdir(VECTORSTORE_DIR):
        # Nếu file đã tồn tại trong vectorstore → xóa chunks cũ trước
        if filename in registry:
            _delete_chunks_by_source(filepath, embeddings)
        vs = load_vectorstore(embeddings)
    else:
        vs = None

    if vs is None:
        vs = Chroma(
            persist_directory=VECTORSTORE_DIR,
            embedding_function=embeddings,
            collection_name=COLLECTION_NAME
        )

    vs.add_documents(chunks)

    # Cập nhật registry
    registry[filename] = current_md5
    _save_registry(registry)

    print(f"   ✅ Đã index {len(chunks)} chunks từ '{filename}'")
    return {"status": "added", "filename": filename, "chunks": len(chunks)}


def _delete_chunks_by_source(filepath: str, embeddings=None):
    """Xóa tất cả chunks có metadata.source = filepath"""
    if embeddings is None:
        embeddings = get_embedding_model()
    vs = load_vectorstore(embeddings)
    results = vs.get(where={"source": filepath})
    ids = results.get("ids", [])
    if ids:
        vs.delete(ids=ids)
        print(f"   🗑️  Đã xóa {len(ids)} chunks cũ của '{filepath}'")


def delete_document(filename: str, embeddings=None) -> dict:
    """
    Xóa 1 file khỏi vectorstore theo tên file.
    Trả về {"status": "deleted"|"not_found"|"error", "filename": ..., "deleted_chunks": ...}
    """
    registry = _load_registry()

    if filename not in registry:
        return {"status": "not_found", "filename": filename, "message": "File chưa được index"}

    if embeddings is None:
        embeddings = get_embedding_model()

    vs = load_vectorstore(embeddings)

    # Tìm theo tên file (source có thể là full path)
    try:
        # Thử tìm theo full path lẫn tên file
        results = vs.get()
        ids_to_delete = []
        for doc_id, meta in zip(results["ids"], results["metadatas"]):
            source = meta.get("source", "")
            if os.path.basename(source) == filename or source == filename:
                ids_to_delete.append(doc_id)

        if ids_to_delete:
            vs.delete(ids=ids_to_delete)

        # Xóa khỏi registry
        del registry[filename]
        _save_registry(registry)

        print(f"   🗑️  Đã xóa {len(ids_to_delete)} chunks của '{filename}'")
        return {"status": "deleted", "filename": filename, "deleted_chunks": len(ids_to_delete)}
    except Exception as e:
        return {"status": "error", "filename": filename, "message": str(e)}


def list_indexed_files() -> list[dict]:
    """Trả về danh sách file đã index"""
    registry = _load_registry()
    result = []
    for filename, md5 in registry.items():
        result.append({"filename": filename, "md5": md5})
    return result


# ── Full reindex ───────────────────────────────────────────────────────────────

def run_indexing(data_dir: str = "data"):
    print("=" * 50)
    print("🚀 BẮT ĐẦU INDEXING PIPELINE")
    print("=" * 50)

    documents = load_documents(data_dir)
    if not documents:
        print("❌ Không load được tài liệu nào!")
        return None

    chunks = chunk_documents(documents)
    embeddings = get_embedding_model()
    vectorstore = build_vectorstore(chunks, embeddings)

    # Cập nhật registry cho tất cả file
    registry = {}
    all_files = []
    for root, dirs, files in os.walk(data_dir):
        for f in files:
            if not f.endswith(".json"):
                all_files.append(os.path.join(root, f))
    for fpath in all_files:
        try:
            registry[os.path.basename(fpath)] = _file_md5(fpath)
        except Exception:
            pass
    _save_registry(registry)

    print("=" * 50)
    print("✅ INDEXING HOÀN TẤT")
    print("=" * 50)
    return vectorstore